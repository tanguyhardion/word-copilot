import os
import re
import requests
import tempfile

from word_copilot.colors import hex_to_rgb
from word_copilot.constants import (
    HEADING_TYPES,
    HEADING_SIZES,
    PAGE_W_PX,
    MARGIN_PX,
    USABLE_W_PX,
)
from word_copilot.icons import (
    download_icon,
    colorize_svg,
    get_svg_aspect_ratio,
    svg_to_png,
    inline_svg_to_tempfile,
    inject_svg_color,
)
from word_copilot.units import px_to_emu, pct_to_px


class WordConnector:
    """
    Renders DSL element lists into a Word document.
    Uses python-docx for flow content and win32com for floating textboxes
    and theme-color access.
    """

    def __init__(self):
        self.word_app = None

    def _connect_com(self):
        import win32com.client, pythoncom

        pythoncom.CoInitialize()
        try:
            self.word_app = win32com.client.GetActiveObject("Word.Application")
        except Exception:
            self.word_app = win32com.client.Dispatch("Word.Application")
            self.word_app.Visible = True
        return self.word_app

    def _prefetch_icons(self, elements, status_cb=None):
        icon_elems = [e for e in elements if e.get("type") == "icon"]
        cache = {}
        for idx, e in enumerate(icon_elems):
            name = e["icon_name"]
            style = e.get("icon_style", "solid")
            key = (name, style)
            if key not in cache:
                if status_cb:
                    status_cb(f"Downloading icon {idx+1}/{len(icon_elems)}: {name}…")
                cache[key] = download_icon(name, style)
        return cache

    def build_and_open(self, pages_data, status_cb=None):
        import pythoncom

        pythoncom.CoInitialize()
        try:
            all_elems = [e for page in pages_data for e in page]
            icon_cache = self._prefetch_icons(all_elems, status_cb)
            if status_cb:
                status_cb("Building document…")
            path = self._build_docx(pages_data, icon_cache, status_cb)
            if status_cb:
                status_cb("Opening in Word…")
            self._connect_com()
            self.word_app.Documents.Open(os.path.abspath(path))
            if status_cb:
                status_cb("✓ Document opened in Word!")
        finally:
            pythoncom.CoUninitialize()

    def insert_at_cursor(self, pages_data, status_cb=None):
        import pythoncom

        pythoncom.CoInitialize()
        try:
            all_elems = [e for page in pages_data for e in page]
            icon_cache = self._prefetch_icons(all_elems, status_cb)
            if status_cb:
                status_cb("Connecting to Word…")
            self._connect_com()
            word = self.word_app
            if word.Documents.Count == 0:
                raise Exception("No document open. Please open one in Word first.")
            doc_com = word.ActiveDocument
            sel = word.Selection
            total_pages = len(pages_data)
            for pi, page_elems in enumerate(pages_data):
                if status_cb:
                    status_cb(f"Inserting page {pi+1}/{total_pages}…")
                if pi > 0:
                    sel.InsertBreak(Type=7)
                for ei, elem in enumerate(page_elems):
                    if status_cb:
                        status_cb(f"Page {pi+1} — element {ei+1}/{len(page_elems)}…")
                    self._insert_elem_com(doc_com, sel, elem, icon_cache)
            if status_cb:
                status_cb(
                    f"✓ Content inserted ({sum(len(p) for p in pages_data)} elements)!"
                )
        finally:
            pythoncom.CoUninitialize()

    def save_docx(self, pages_data, status_cb=None):
        all_elems = [e for page in pages_data for e in page]
        icon_cache = self._prefetch_icons(all_elems, status_cb)
        return self._build_docx(pages_data, icon_cache, status_cb)

    def _build_docx(self, pages_data, icon_cache, status_cb=None):
        from docx import Document

        doc = Document()
        page_cfg = self._extract_page_cfg(pages_data)
        self._apply_page_setup(doc, page_cfg)
        usable_w_px = page_cfg.get("usable_w_px", USABLE_W_PX)

        total_pages = len(pages_data)
        for pi, page_elems in enumerate(pages_data):
            if pi > 0:
                pb_para = doc.add_paragraph()
                pb_run = pb_para.add_run()
                pb_run.add_break(
                    __import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE
                )
            for ei, elem in enumerate(page_elems):
                if status_cb:
                    status_cb(
                        f"Page {pi+1}/{total_pages} — element {ei+1}/{len(page_elems)}…"
                    )
                self._render_elem(doc, elem, icon_cache, usable_w_px)

        out_dir = os.path.join(tempfile.gettempdir(), "word_dsl_output")
        os.makedirs(out_dir, exist_ok=True)
        path = os.path.join(out_dir, "word_dsl_output.docx")
        doc.save(path)
        print(f"[docx] Saved to {path}")
        return path

    def _extract_page_cfg(self, pages_data):
        cfg = {
            "size": "letter",
            "orientation": "portrait",
            "margin_top": MARGIN_PX,
            "margin_right": MARGIN_PX,
            "margin_bottom": MARGIN_PX,
            "margin_left": MARGIN_PX,
        }
        for page_elems in pages_data:
            for elem in page_elems:
                if elem.get("type") == "page":
                    cfg.update({k: v for k, v in elem.items() if k != "type"})
                    break
            break
        w_px = PAGE_W_PX
        if cfg["size"] == "a4":
            w_px = 794
        if cfg["orientation"] == "landscape":
            w_px, _ = _, w_px
        cfg["usable_w_px"] = w_px - cfg["margin_left"] - cfg["margin_right"]
        return cfg

    def _apply_page_setup(self, doc, cfg):
        section = doc.sections[0]
        size = cfg.get("size", "letter")
        orientation = cfg.get("orientation", "portrait")
        if size == "a4":
            w_emu = px_to_emu(794)
            h_emu = px_to_emu(1123)
        else:
            w_emu = px_to_emu(816)
            h_emu = px_to_emu(1056)
        if orientation == "landscape":
            w_emu, h_emu = h_emu, w_emu
        section.page_width = w_emu
        section.page_height = h_emu
        section.top_margin = px_to_emu(cfg.get("margin_top", MARGIN_PX))
        section.right_margin = px_to_emu(cfg.get("margin_right", MARGIN_PX))
        section.bottom_margin = px_to_emu(cfg.get("margin_bottom", MARGIN_PX))
        section.left_margin = px_to_emu(cfg.get("margin_left", MARGIN_PX))

    def _render_elem(self, doc, elem, icon_cache, usable_w_px):
        t = elem.get("type")
        if t == "page":
            return
        if t in HEADING_TYPES:
            self._render_heading(doc, elem)
        elif t == "p":
            self._render_paragraph(doc, elem)
        elif t in ("ul", "ol"):
            self._render_list(doc, elem)
        elif t == "item":
            self._render_list_item(doc, elem, list_style="ul", level=0)
        elif t == "hr":
            self._render_hr(doc, elem)
        elif t == "br":
            doc.add_paragraph()
        elif t == "pagebreak":
            from docx.enum.text import WD_BREAK

            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)
        elif t == "image":
            self._render_image(doc, elem, usable_w_px)
        elif t == "icon":
            self._render_icon(doc, elem, icon_cache, usable_w_px)
        elif t == "svg":
            self._render_svg(doc, elem, usable_w_px)
        elif t == "table":
            self._render_table(doc, elem, usable_w_px)
        elif t == "textbox":
            self._render_textbox_docx(doc, elem)

    def _render_heading(self, doc, elem):
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        level = elem.get("level", 1)
        style_name = f"Heading {level}"
        para = doc.add_paragraph(style=style_name)
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if "align" in elem:
            para.alignment = align_map.get(elem["align"], WD_ALIGN_PARAGRAPH.LEFT)
        if "spacing_before" in elem:
            para.paragraph_format.space_before = Pt(elem["spacing_before"])
        if "spacing_after" in elem:
            para.paragraph_format.space_after = Pt(elem["spacing_after"])
        rich = elem.get("rich_text", [])
        if not rich:
            return
        default_size = elem.get("size", HEADING_SIZES.get(f"h{level}", 14))
        default_color = elem.get("color")
        default_bold = elem.get("bold", True)
        default_font = elem.get("font")
        for seg in rich:
            run = para.add_run(seg.get("text", ""))
            run.font.size = Pt(seg.get("size", default_size))
            run.font.bold = seg.get("bold", default_bold)
            run.font.italic = seg.get("italic", False)
            if seg.get("underline"):
                run.font.underline = True
            if seg.get("font") or default_font:
                run.font.name = seg.get("font", default_font)
            color = seg.get("color") or default_color
            if color:
                r, g, b = hex_to_rgb(color)
                run.font.color.rgb = RGBColor(r, g, b)

    def _render_paragraph(self, doc, elem):
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        para = doc.add_paragraph()
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if "align" in elem:
            para.alignment = align_map.get(elem["align"], WD_ALIGN_PARAGRAPH.LEFT)
        if "spacing_before" in elem:
            para.paragraph_format.space_before = Pt(elem["spacing_before"])
        if "spacing_after" in elem:
            para.paragraph_format.space_after = Pt(elem["spacing_after"])
        if "line_height" in elem:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            pPr = para._p.get_or_add_pPr()
            lnSpc = OxmlElement("w:lnSpc")
            spcPct = OxmlElement("w:pct")
            spcPct.set(qn("w:val"), str(int(elem["line_height"] * 100)))
            lnSpc.append(spcPct)
            pPr.append(lnSpc)
        rich = elem.get("rich_text", [])
        default_size = elem.get("size", 11)
        default_color = elem.get("color")
        default_bold = elem.get("bold", False)
        default_italic = elem.get("italic", False)
        default_font = elem.get("font")
        for seg in rich:
            run = para.add_run(seg.get("text", ""))
            run.font.size = Pt(seg.get("size", default_size))
            run.font.bold = seg.get("bold", default_bold)
            run.font.italic = seg.get("italic", default_italic)
            if seg.get("underline"):
                run.font.underline = True
            if seg.get("font") or default_font:
                run.font.name = seg.get("font", default_font)
            color = seg.get("color") or default_color
            if color:
                r, g, b = hex_to_rgb(color)
                run.font.color.rgb = RGBColor(r, g, b)

    def _render_list(self, doc, list_elem):
        list_type = list_elem.get("type", "ul")
        fields = list_elem.get("fields", {})
        indent = int(fields.get("indent", 1))
        items = list_elem.get("items", [])
        for item in items:
            self._render_list_item(doc, item, list_style=list_type, level=indent - 1)

    def _render_list_item(self, doc, item, list_style="ul", level=0):
        from docx.shared import Pt, RGBColor

        style_name = "List Bullet" if list_style == "ul" else "List Number"
        if level > 0:
            style_name += f" {level+1}" if level < 3 else " 3"
        try:
            para = doc.add_paragraph(style=style_name)
        except Exception:
            para = doc.add_paragraph(style="List Bullet")
        rich = item.get("rich_text", [])
        default_size = item.get("size", 11)
        default_color = item.get("color")
        default_bold = item.get("bold", False)
        default_italic = item.get("italic", False)
        default_font = item.get("font")
        for seg in rich:
            run = para.add_run(seg.get("text", ""))
            run.font.size = Pt(seg.get("size", default_size))
            run.font.bold = seg.get("bold", default_bold)
            run.font.italic = seg.get("italic", default_italic)
            if seg.get("underline"):
                run.font.underline = True
            if seg.get("font") or default_font:
                run.font.name = seg.get("font", default_font)
            color = seg.get("color") or default_color
            if color:
                r, g, b = hex_to_rgb(color)
                run.font.color.rgb = RGBColor(r, g, b)

    def _render_hr(self, doc, elem):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        color_hex = elem.get("color", "#AAAAAA")
        if color_hex.startswith("#"):
            color_hex = color_hex[1:]
        weight_pt = elem.get("weight", 1.0)
        size_val = str(max(1, int(weight_pt * 8)))
        para = doc.add_paragraph()
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), size_val)
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color_hex.upper())
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _render_image(self, doc, elem, usable_w_px):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        url = elem.get("url", "")
        desired_w = elem.get("width", usable_w_px)
        desired_h = elem.get("height", None)
        align = elem.get("align", "left")
        try:
            headers = {"User-Agent": "Mozilla/5.0"}
            resp = requests.get(url, timeout=15, verify=False, headers=headers)
            resp.raise_for_status()
            from urllib.parse import urlparse

            filename = os.path.basename(urlparse(url).path) or "image.jpg"
            filename = re.sub(r'[<>:"/\\|?*]', "_", filename)
            tmp_path = os.path.join(tempfile.gettempdir(), filename)
            with open(tmp_path, "wb") as f:
                f.write(resp.content)
            para = doc.add_paragraph()
            align_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }
            para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
            run = para.add_run()
            w_emu = px_to_emu(desired_w)
            if desired_h:
                run.add_picture(tmp_path, width=w_emu, height=px_to_emu(desired_h))
            else:
                run.add_picture(tmp_path, width=w_emu)
        except Exception as e:
            print(f"[image] Failed: {e}")
            p = doc.add_paragraph(f"[Image: {url}]")
            p.runs[0].italic = True

    def _render_icon(self, doc, elem, icon_cache, usable_w_px):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        name = elem.get("icon_name", "question")
        style = elem.get("icon_style", "solid")
        color = elem.get("icon_color")
        width = elem.get("width", 32)
        align = elem.get("align", "left")
        svg_path = icon_cache.get((name, style))
        if not svg_path or not os.path.exists(svg_path):
            p = doc.add_paragraph(f"[Icon: {name}]")
            p.runs[0].italic = True
            return
        working_path = colorize_svg(svg_path, color) if color else svg_path
        aspect = get_svg_aspect_ratio(working_path)
        if aspect <= 0:
            aspect = 1.0
        height = width / aspect
        png_path = svg_to_png(
            working_path,
            width_px=max(1, int(width * 2)),
            height_px=max(1, int(height * 2)),
        )
        insert_path = (
            png_path if png_path and os.path.exists(png_path) else working_path
        )
        try:
            para = doc.add_paragraph()
            align_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }
            para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
            run = para.add_run()
            run.add_picture(os.path.abspath(insert_path), width=px_to_emu(width))
        except Exception as e:
            print(f"[icon] Insert failed for '{name}': {e}")
            doc.add_paragraph(f"[Icon: {name}]").runs[0].italic = True

    def _render_svg(self, doc, elem, usable_w_px):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        svg_markup = elem.get("svg_markup", "")
        color = elem.get("svg_color")
        width = elem.get("width", 100)
        align = elem.get("align", "left")
        if color:
            svg_markup = inject_svg_color(svg_markup, color)
        svg_path = inline_svg_to_tempfile(svg_markup)
        if not svg_path:
            return
        aspect = get_svg_aspect_ratio(svg_path)
        if aspect <= 0:
            aspect = 1.0
        height = width / aspect
        png_path = svg_to_png(
            svg_path, width_px=max(1, int(width * 2)), height_px=max(1, int(height * 2))
        )
        insert_path = png_path if png_path and os.path.exists(png_path) else svg_path
        try:
            para = doc.add_paragraph()
            align_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }
            para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
            run = para.add_run()
            run.add_picture(os.path.abspath(insert_path), width=px_to_emu(width))
        except Exception as e:
            print(f"[svg] Insert failed: {e}")

    def _render_table(self, doc, elem, usable_w_px):
        from docx.shared import Pt, RGBColor

        td = elem.get("table", {})
        style = elem.get("table_style", {})
        rows = td.get("rows", 0)
        cols = td.get("cols", 0)
        if rows == 0 or cols == 0:
            return
        raw_w = elem.get("width_spec", "100%")
        total_w_px = pct_to_px(raw_w, usable_w_px)
        col_specs = td.get("col_widths_spec", [])
        if col_specs:
            col_widths_px = [pct_to_px(s, total_w_px) for s in col_specs[:cols]]
            while len(col_widths_px) < cols:
                col_widths_px.append(total_w_px / cols)
        else:
            col_widths_px = [total_w_px / cols] * cols
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        for ci, cw in enumerate(col_widths_px):
            for row in table.rows:
                row.cells[ci].width = px_to_emu(cw)
        content = td.get("content", [])
        header_row = td.get("header_row", False)
        for ri in range(min(rows, len(content))):
            is_header = ri == 0 and header_row
            for ci in range(min(cols, len(content[ri]))):
                cell = table.cell(ri, ci)
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(str(content[ri][ci]))
                if style.get("font"):
                    run.font.name = style["font"]
                if style.get("font_size"):
                    run.font.size = Pt(style["font_size"])
                if is_header:
                    run.font.bold = style.get("header_bold", True)
                    text_color = style.get("header_text_color") or style.get(
                        "text_color"
                    )
                else:
                    text_color = style.get("text_color")
                if text_color:
                    r, g, b = hex_to_rgb(text_color)
                    run.font.color.rgb = RGBColor(r, g, b)
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                align_map = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                }
                para.alignment = align_map.get(
                    style.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT
                )
                fill_color = None
                if is_header:
                    fill_color = style.get("header_fill")
                elif ri % 2 == 0:
                    fill_color = style.get("row_fill")
                else:
                    fill_color = style.get("alt_row_fill") or style.get("row_fill")
                if fill_color:
                    self._set_cell_fill(cell, fill_color)
                border_color = style.get("border_color")
                border_weight = style.get("border_weight", 1.0)
                if border_color:
                    self._set_cell_borders(cell, border_color, border_weight)

    def _set_cell_fill(self, cell, hex_color):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        hex_color = hex_color.lstrip("#")
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color.upper())
        tcPr.append(shd)

    def _set_cell_borders(self, cell, hex_color, weight_pt=1.0):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        hex_color = hex_color.lstrip("#").upper()
        size_val = str(max(1, int(weight_pt * 8)))
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size_val)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), hex_color)
            tcBorders.append(el)
        tcPr.append(tcBorders)

    def _render_textbox_docx(self, doc, elem):
        import lxml.etree as etree

        left = px_to_emu(elem.get("left", 0))
        top = px_to_emu(elem.get("top", 0))
        width = px_to_emu(elem.get("width", 150))
        height = px_to_emu(elem.get("height", 60))
        fill_hex = elem.get("color", "#CCCCCC").lstrip("#")
        rich = elem.get("rich_text", [])

        body_xml_parts = []
        for seg in rich:
            text = seg.get("text", "")
            size_half = int(seg.get("size", 11) * 2)
            bold_val = "1" if seg.get("bold") else "0"
            ital_val = "1" if seg.get("italic") else "0"
            color_hex = seg.get("color", "#000000").lstrip("#")
            font_name = seg.get("font", "Calibri")
            body_xml_parts.append(
                f'<a:r xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
                f'<a:rPr lang="en-US" sz="{size_half}" b="{bold_val}" i="{ital_val}" '
                f'dirty="0">'
                f'<a:solidFill><a:srgbClr val="{color_hex.upper()}"/></a:solidFill>'
                f'<a:latin typeface="{font_name}"/>'
                f"</a:rPr>"
                f"<a:t>{text}</a:t>"
                f"</a:r>"
            )

        body_xml = "".join(body_xml_parts)
        h_align = elem.get("halign", "left")
        algn_map = {"left": "l", "center": "ctr", "right": "r", "justify": "just"}
        algn = algn_map.get(h_align, "l")
        v_align = elem.get("valign", "top")
        anchor_map = {"top": "t", "middle": "ctr", "bottom": "b"}
        anchor = anchor_map.get(v_align, "t")
        pad_l = px_to_emu(elem.get("padding_left", 7))
        pad_r = px_to_emu(elem.get("padding_right", 7))
        pad_t = px_to_emu(elem.get("padding_top", 3))
        pad_b = px_to_emu(elem.get("padding_bottom", 3))
        outline = elem.get("outline")
        if outline:
            ol_hex = outline["color"].lstrip("#")
            ol_w = int(outline.get("weight", 1) * 12700)
            ln_xml = (
                f'<a:ln w="{ol_w}" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
                f'<a:solidFill><a:srgbClr val="{ol_hex.upper()}"/></a:solidFill>'
                f"</a:ln>"
            )
        else:
            ln_xml = '<a:ln xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><a:noFill/></a:ln>'

        drawing_xml = f"""
<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:r>
    <w:rPr/>
    <mc:AlternateContent
        xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">
      <mc:Choice Requires="wps">
        <w:drawing>
          <wp:anchor distT="0" distB="0" distL="114300" distR="114300"
              simplePos="0" relativeHeight="251658240" behindDoc="0"
              locked="0" layoutInCell="1" allowOverlap="1"
              xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="page">
              <wp:posOffset>{left}</wp:posOffset>
            </wp:positionH>
            <wp:positionV relativeFrom="page">
              <wp:posOffset>{top}</wp:posOffset>
            </wp:positionV>
            <wp:extent cx="{width}" cy="{height}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapNone/>
            <wp:docPr id="1" name="TextBox"/>
            <wp:cNvGraphicFramePr/>
            <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
              <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                <wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                  <wps:cNvSpPr><a:spLocks noChangeArrowheads="1"/></wps:cNvSpPr>
                  <wps:spPr>
                    <a:xfrm>
                      <a:off x="0" y="0"/>
                      <a:ext cx="{width}" cy="{height}"/>
                    </a:xfrm>
                    <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                    <a:solidFill>
                      <a:srgbClr val="{fill_hex.upper()}"/>
                    </a:solidFill>
                    {ln_xml}
                  </wps:spPr>
                  <wps:txbx>
                    <w:txbxContent>
                      <w:p>
                        <w:pPr>
                          <w:jc w:val="{algn}"/>
                        </w:pPr>
                        <w:r>
                          <w:t xml:space="preserve"> </w:t>
                        </w:r>
                      </w:p>
                    </w:txbxContent>
                  </wps:txbx>
                  <wps:bodyPr anchor="{anchor}"
                      lIns="{pad_l}" rIns="{pad_r}"
                      tIns="{pad_t}" bIns="{pad_b}">
                    <a:spAutoFit/>
                  </wps:bodyPr>
                </wps:wsp>
              </a:graphicData>
            </a:graphic>
          </wp:anchor>
        </w:drawing>
      </mc:Choice>
    </mc:AlternateContent>
  </w:r>
</w:p>"""

        try:
            drawing_elem = etree.fromstring(drawing_xml.strip())
            ns = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
                "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            }
            txbx_content = drawing_elem.find(".//wps:txbx/w:txbxContent", ns)
            if txbx_content is not None and body_xml:
                for child in list(txbx_content):
                    txbx_content.remove(child)
                inner_para = etree.fromstring(
                    f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                    f'     xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
                    f'<w:pPr><w:jc w:val="{algn}"/></w:pPr>' + body_xml + f"</w:p>"
                )
                txbx_content.append(inner_para)
            doc.element.body.append(drawing_elem)
        except Exception as e:
            print(f"[textbox] XML injection failed: {e}")
            p = doc.add_paragraph()
            for seg in rich:
                run = p.add_run(seg.get("text", ""))
                run.font.bold = seg.get("bold", False)

    def _insert_elem_com(self, doc_com, sel, elem, icon_cache):
        t = elem.get("type")
        if t == "page":
            return
        if t in HEADING_TYPES:
            self._com_insert_heading(sel, elem)
        elif t == "p":
            self._com_insert_paragraph(sel, elem)
        elif t in ("ul", "ol"):
            self._com_insert_list(sel, elem)
        elif t == "hr":
            self._com_insert_hr(sel, elem)
        elif t == "br":
            sel.TypeParagraph()
        elif t == "pagebreak":
            sel.InsertBreak(Type=7)
        elif t == "image":
            self._com_insert_image(sel, elem)
        elif t == "icon":
            self._com_insert_icon(sel, elem, icon_cache)
        elif t in ("svg", "textbox", "table"):
            self._com_insert_via_docx(doc_com, sel, elem, icon_cache)

    def _com_insert_heading(self, sel, elem):
        level = elem.get("level", 1)
        style_name = f"Heading {level}"
        rich = elem.get("rich_text", [])
        text = "".join(s.get("text", "") for s in rich)
        try:
            sel.Style = sel.Document.Styles(style_name)
        except Exception:
            pass
        sel.TypeText(text)
        sel.TypeParagraph()

    def _com_insert_paragraph(self, sel, elem):
        rich = elem.get("rich_text", [])
        try:
            sel.Style = sel.Document.Styles("Normal")
        except Exception:
            pass
        for seg in rich:
            sel.Font.Bold = seg.get("bold", False)
            sel.Font.Italic = seg.get("italic", False)
            sel.Font.Underline = 1 if seg.get("underline") else 0
            if seg.get("size"):
                sel.Font.Size = seg["size"]
            if seg.get("font"):
                sel.Font.Name = seg["font"]
            if seg.get("color"):
                r, g, b = hex_to_rgb(seg["color"])
                sel.Font.Color = r + (g << 8) + (b << 16)
            sel.TypeText(seg.get("text", ""))
        sel.Font.Reset()
        sel.TypeParagraph()

    def _com_insert_list(self, sel, list_elem):
        list_type = list_elem.get("type", "ul")
        items = list_elem.get("items", [])
        list_id = 1 if list_type == "ul" else 2
        for item in items:
            rich = item.get("rich_text", [])
            text = "".join(s.get("text", "") for s in rich)
            try:
                sel.Range.ListFormat.ApplyListTemplate(
                    sel.Document.ListTemplates(list_id)
                )
            except Exception:
                pass
            sel.TypeText(text)
            sel.TypeParagraph()

    def _com_insert_hr(self, sel, elem):
        sel.TypeParagraph()
        try:
            color_hex = elem.get("color", "#AAAAAA")
            r, g, b = hex_to_rgb(color_hex)
            rgb_val = r + (g << 8) + (b << 16)
            weight_pt = elem.get("weight", 1.0)
            para = sel.Paragraphs(1)
            borders = para.Borders
            bottom = borders(3)
            bottom.LineStyle = 1
            bottom.LineWidth = max(1, int(weight_pt * 8))
            bottom.Color = rgb_val
        except Exception as e:
            print(f"[hr] COM border warning: {e}")

    def _com_insert_image(self, sel, elem):
        url = elem.get("url", "")
        try:
            headers = {"User-Agent": "Mozilla/5.0"}
            resp = requests.get(url, timeout=15, verify=False, headers=headers)
            resp.raise_for_status()
            from urllib.parse import urlparse

            filename = os.path.basename(urlparse(url).path) or "image.jpg"
            filename = re.sub(r'[<>:"/\\|?*]', "_", filename)
            tmp = os.path.join(tempfile.gettempdir(), filename)
            with open(tmp, "wb") as f:
                f.write(resp.content)
            width_px = elem.get("width", 200)
            height_px = elem.get("height", None)
            w_pt = width_px * 0.75
            pic = sel.InlineShapes.AddPicture(
                FileName=os.path.abspath(tmp),
                LinkToFile=False,
                SaveWithDocument=True,
                Range=sel.Range,
            )
            pic.LockAspectRatio = True
            pic.Width = w_pt
            if height_px:
                pic.LockAspectRatio = False
                pic.Height = height_px * 0.75
        except Exception as e:
            print(f"[image] COM insert failed: {e}")
            sel.TypeText(f"[Image: {url}]")
            sel.TypeParagraph()

    def _com_insert_icon(self, sel, elem, icon_cache):
        name = elem.get("icon_name", "question")
        style = elem.get("icon_style", "solid")
        color = elem.get("icon_color")
        width = elem.get("width", 32)
        svg_path = icon_cache.get((name, style))
        if not svg_path or not os.path.exists(svg_path):
            sel.TypeText(f"[Icon: {name}]")
            sel.TypeParagraph()
            return
        working = colorize_svg(svg_path, color) if color else svg_path
        aspect = get_svg_aspect_ratio(working)
        if aspect <= 0:
            aspect = 1.0
        height = width / aspect
        png_path = svg_to_png(
            working, width_px=max(1, int(width * 2)), height_px=max(1, int(height * 2))
        )
        insert_path = png_path if png_path and os.path.exists(png_path) else working
        try:
            pic = sel.InlineShapes.AddPicture(
                FileName=os.path.abspath(insert_path),
                LinkToFile=False,
                SaveWithDocument=True,
                Range=sel.Range,
            )
            pic.LockAspectRatio = True
            pic.Width = width * 0.75
        except Exception as e:
            print(f"[icon] COM insert failed: {e}")
            sel.TypeText(f"[Icon: {name}]")

    def _com_insert_via_docx(self, doc_com, sel, elem, icon_cache):
        from docx import Document

        mini_doc = Document()
        self._render_elem(mini_doc, elem, icon_cache, USABLE_W_PX)
        out_dir = os.path.join(tempfile.gettempdir(), "word_dsl_mini")
        os.makedirs(out_dir, exist_ok=True)
        path = os.path.join(out_dir, f"mini_{elem.get('type','elem')}.docx")
        mini_doc.save(path)
        try:
            sel.InsertFile(
                FileName=os.path.abspath(path),
                ConfirmConversions=False,
                Link=False,
                Attachment=False,
            )
        except Exception as e:
            print(f"[com] InsertFile failed: {e}")
