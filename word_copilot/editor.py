import copy
import os
from word_copilot.constants import USABLE_W_PX
from word_copilot.connector import WordConnector
from word_copilot.extractor import WordExtractor


class WordEditor:
    """
    Applies a list of edit operations (from parse_edit_dsl) to a *snapshot*
    of the currently active Word document, then reopens the edited copy in
    Word (replacing the old window).

    Targeting rule (must match WordExtractor's id assignment exactly):
    id = 0-based index of a <w:p> or <w:tbl> element among the document
    body's direct children, in document order.

    New content for replace/insert_* ops is produced by rendering each
    element through the *existing* WordConnector._render_elem into a
    disposable scratch Document(), then deep-copying the resulting
    <w:p>/<w:tbl> XML nodes into the target document. This reuses every
    existing renderer (headings, tables, icons, svg, textboxes, lists...)
    without duplicating any rendering logic.
    """

    def __init__(self):
        self.extractor = WordExtractor()
        self.connector = WordConnector()

    # ── public entry point ────────────────────────────────────────────────────
    def apply(self, edit_ops, status_cb=None):
        import pythoncom

        pythoncom.CoInitialize()
        try:
            if status_cb:
                status_cb("Locating active document…")
            path = self.extractor._get_active_doc_path_for_inplace_edit(status_cb)

            all_new_elems = []
            for op in edit_ops:
                all_new_elems.extend(op.get("elements", []))
            icon_cache = self.connector._prefetch_icons(all_new_elems, status_cb)

            if status_cb:
                status_cb("Applying edits…")

            scroll_target_id = None
            for op in edit_ops:
                tid = op.get("target_id")
                if tid is not None:
                    if scroll_target_id is None or tid < scroll_target_id:
                        scroll_target_id = tid

            edited_path = self._apply_ops_to_docx(path, edit_ops, icon_cache, status_cb)

            if status_cb:
                status_cb("Reopening document in Word…")
            self._reopen_in_word(edited_path, scroll_target_id, status_cb)

            if status_cb:
                status_cb("✓ Edits applied!")
            return edited_path
        finally:
            pythoncom.CoUninitialize()

    # ── core XML surgery ──────────────────────────────────────────────────────
    def _apply_ops_to_docx(self, path, edit_ops, icon_cache, status_cb=None):
        from docx import Document

        doc = Document(path)
        body = doc.element.body

        def targetable_children():
            return [
                c
                for c in body
                if (c.tag.split("}")[-1] if "}" in c.tag else c.tag) in ("p", "tbl")
            ]

        children = targetable_children()

        total = len(edit_ops)
        for oi, op in enumerate(edit_ops):
            if status_cb:
                status_cb(f"Edit {oi+1}/{total}: {op['op']}…")

            if op["op"] == "delete":
                el = self._resolve_target(children, op["target_id"])
                if el is None:
                    continue
                el.getparent().remove(el)

            elif op["op"] == "replace":
                el = self._resolve_target(children, op["target_id"])
                if el is None:
                    continue
                new_nodes = self._render_elements_to_nodes(op["elements"], icon_cache)
                anchor = el
                for node in new_nodes:
                    anchor.addnext(node)
                    anchor = node
                el.getparent().remove(el)

            elif op["op"] in ("insert_before", "insert_after"):
                el = self._resolve_target(children, op["target_id"])
                if el is None:
                    continue
                new_nodes = self._render_elements_to_nodes(op["elements"], icon_cache)
                if op["op"] == "insert_after":
                    anchor = el
                    for node in new_nodes:
                        anchor.addnext(node)
                        anchor = node
                else:  # insert_before
                    for node in new_nodes:
                        el.addprevious(node)

            elif op["op"] == "insert_at":
                new_nodes = self._render_elements_to_nodes(op["elements"], icon_cache)
                if op.get("position") == "start":
                    children_now = targetable_children()
                    if children_now:
                        first = children_now[0]
                        for node in new_nodes:
                            first.addprevious(node)
                    else:
                        for node in new_nodes:
                            body.insert(0, node)
                else:  # "end" (default)
                    sect_pr = body.find(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr"
                    )
                    for node in new_nodes:
                        if sect_pr is not None:
                            sect_pr.addprevious(node)
                        else:
                            body.append(node)

            else:
                print(f"[edit] unknown op '{op['op']}' — skipped")

        doc.save(path)
        return path

    def _resolve_target(self, children, target_id):
        if target_id < 0 or target_id >= len(children):
            print(
                f"[edit] target id={target_id} out of range "
                f"(document has {len(children)} addressable elements) — skipped"
            )
            return None
        el = children[target_id]
        if el.getparent() is None:
            print(
                f"[edit] target id={target_id} was already removed/replaced "
                f"by an earlier operation in this batch — skipped"
            )
            return None
        return el

    def _render_elements_to_nodes(self, elements, icon_cache):
        from docx import Document

        nodes = []
        for elem in elements:
            scratch = Document()
            if elem.get("type") == "page":
                continue
            self.connector._render_elem(scratch, elem, icon_cache, USABLE_W_PX)
            scratch_children = [
                c
                for c in scratch.element.body
                if (c.tag.split("}")[-1] if "}" in c.tag else c.tag) in ("p", "tbl")
            ]
            for c in scratch_children:
                nodes.append(copy.deepcopy(c))
        return nodes

    # ── reopen edited doc in Word ─────────────────────────────────────────────
    def _reopen_in_word(self, edited_path, scroll_target_id=None, status_cb=None):
        import win32com.client, time

        try:
            word = win32com.client.GetActiveObject("Word.Application")
        except Exception:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = True

        doc = word.Documents.Open(os.path.abspath(edited_path))
        word.Visible = True
        word.Activate()

        if scroll_target_id is None:
            return
        try:
            if status_cb:
                status_cb("Scrolling to edited element…")

            import time as _time

            deadline = _time.time() + 5.0
            while _time.time() < deadline:
                try:
                    _ = doc.ActiveWindow.Panes(1).Pages.Count
                    break
                except Exception:
                    _time.sleep(0.2)

            win = doc.Windows(1)
            word.Visible = True
            doc.Activate()
            win.Activate()

            from docx import Document as _Doc
            from docx.oxml.ns import qn as _qn

            _d = _Doc(edited_path)
            _body = _d.element.body

            def _count_paragraphs_in_element(el):
                wp_tag = _qn("w:p")
                return len(el.findall(f".//{wp_tag}"))

            para_idx = 0
            body_idx = -1
            target_para_idx = None

            for child in _body:
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if tag not in ("p", "tbl"):
                    continue
                body_idx += 1

                if tag == "p":
                    para_idx += 1
                    if body_idx == scroll_target_id:
                        target_para_idx = para_idx
                        break
                elif tag == "tbl":
                    table_para_count = _count_paragraphs_in_element(child)
                    para_idx += table_para_count
                    if body_idx == scroll_target_id:
                        target_para_idx = (
                            (para_idx - table_para_count + 1)
                            if table_para_count > 0
                            else (para_idx + 1)
                        )
                        break

            para_count = doc.Paragraphs.Count
            goto_idx = max(1, min(target_para_idx or 1, para_count))

            rng = doc.Paragraphs(goto_idx).Range
            rng.Select()
            _time.sleep(0.15)

            target_page = rng.Information(3)

            page_rng = doc.GoTo(What=1, Which=1, Count=target_page)
            page_rng.Select()
            win.ScrollIntoView(page_rng)

            rng.Select()
            win.ScrollIntoView(rng)

        except Exception as e:
            print(f"[scroll] Could not scroll to edit target: {e}")
